from __future__ import annotations

import csv
import io
import hashlib
import json
import sqlite3
import time
from contextlib import contextmanager
from dataclasses import replace
from datetime import date, datetime, timedelta, timezone
from email.message import EmailMessage
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from msgforge import Message
from openpyxl import Workbook
from pypdf import PdfWriter
from pypdf.generic import DictionaryObject, DecodedStreamObject, NameObject
import pytest

from portfolio_assistant import APPLICATION_ID
from portfolio_assistant.api import create_app
from portfolio_assistant.llm import LlmContractError


def upload(client: TestClient, project_id: str, name: str, data: bytes, **fields):
    return client.post(
        f"/api/projects/{project_id}/sources",
        files={"file": (name, data, "application/octet-stream")},
        data={key: str(value).lower() if isinstance(value, bool) else value for key, value in fields.items()},
    )


def process(client: TestClient, source_id: int):
    response = client.post(f"/api/sources/{source_id}/retry")
    assert response.status_code == 200, response.text
    return response.json()


def make_eml() -> bytes:
    message = EmailMessage()
    message["Subject"] = "Fictional Atlas design decision"
    message["From"] = "avery.chen@fictional.example"
    message["To"] = "casey.morgan@fictional.example"
    message["Date"] = "Tue, 11 Aug 2026 09:30:00 -0400"
    message["Message-ID"] = "<atlas-decision-001@fictional.example>"
    message.set_content("The fictional Atlas team approved the blue release plan.\nACTION: Publish the release checklist | OWNER: Me | DUE: 2026-08-20")
    message.add_attachment(b"Attachment evidence says the validation window opens Friday.", maintype="text", subtype="plain", filename="fictional-evidence.txt")
    return message.as_bytes()


def make_msg(tmp_path: Path) -> bytes:
    path = tmp_path / "fictional-message.msg"
    message = Message(
        subject="Fictional Atlas readiness",
        text_body="The fictional readiness review is scheduled for Thursday.",
        to=[("casey.morgan@fictional.example", "Casey Morgan")],
        sender=("avery.chen@fictional.example", "Avery Chen"),
        sent=True,
    )
    message.attach_bytes("readiness-note.txt", b"Fictional attachment confirms the test window.")
    message.save(path)
    return path.read_bytes()


def make_docx(tmp_path: Path) -> bytes:
    path = tmp_path / "fictional-plan.docx"
    document = Document()
    document.add_heading("Fictional Atlas Plan", 1)
    document.add_paragraph("The rollout rehearsal is approved for the certification environment.")
    document.save(path)
    return path.read_bytes()


def make_pdf(tmp_path: Path, with_text: bool = True) -> bytes:
    path = tmp_path / ("text.pdf" if with_text else "scanned.pdf")
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    if with_text:
        font = DictionaryObject({
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        })
        font_ref = writer._add_object(font)
        page[NameObject("/Resources")] = DictionaryObject({
            NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})
        })
        stream = DecodedStreamObject()
        stream.set_data(b"BT /F1 12 Tf 72 720 Td (Fictional Atlas PDF evidence approves the release.) Tj ET")
        page[NameObject("/Contents")] = writer._add_object(stream)
    with path.open("wb") as handle:
        writer.write(handle)
    return path.read_bytes()


def snow_csv(comments: str, *, updated: str = "2026-08-12 10:00:00", state: str = "In Progress", priority: str = "2") -> bytes:
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=[
        "Number", "Short description", "Assignment group", "Updated",
        "Comments and Work notes", "State", "Priority",
    ])
    writer.writeheader()
    writer.writerow({
        "Number": "REQ0099001", "Short description": "Fictional SNOW Migration",
        "Assignment group": "Fictional Clinical Apps", "Updated": updated,
        "Comments and Work notes": comments, "State": state, "Priority": priority,
    })
    return output.getvalue().encode("utf-8")


def snow_xlsx(comments: str) -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["Number", "Short description", "Assignment group", "Updated", "Comments and Work notes"])
    sheet.append(["REQ0099002", "Fictional XLSX Intake", "Fictional Infrastructure", datetime(2026, 8, 12, 11, 0), comments])
    stream = io.BytesIO()
    workbook.save(stream)
    return stream.getvalue()


def test_environment_loopback_configuration_and_restart(settings, client: TestClient):
    health = client.get("/api/health").json()
    assert health["application"] == APPLICATION_ID
    assert health["retrieval_mode"] == "fts5"
    rejected = client.get("http://evil.example/api/health", headers={"host": "evil.example"})
    assert rejected.status_code == 400
    origin = client.get("/api/health", headers={"origin": "https://evil.example"})
    assert origin.status_code == 403
    other_loopback_port = client.get("/api/health", headers={"origin": "http://127.0.0.1:9999"})
    assert other_loopback_port.status_code == 403
    same_origin = client.get("/api/health", headers={"origin": "http://127.0.0.1:8765"})
    assert same_origin.status_code == 200
    missing_origin = client.post(
        "/api/groups", json={"name": "Rejected form"}, headers={"X-Requested-With": ""}
    )
    assert missing_origin.status_code == 403
    same_origin_write = client.post(
        "/api/groups", json={"name": "Fictional Same Origin"},
        headers={"Origin": "http://127.0.0.1:8765", "X-Requested-With": ""},
    )
    assert same_origin_write.status_code == 201
    oversized_request = client.post(
        "/api/groups", content=b"", headers={"Content-Length": str(3 * 1024 * 1024)}
    )
    assert oversized_request.status_code == 413
    streamed_oversized_request = client.post(
        "/api/groups",
        content=iter([b'{"name":"', *(b"x" * (800 * 1024) for _ in range(3)), b'"}']),
        headers={"Content-Type": "application/json"},
    )
    assert streamed_oversized_request.status_code == 413, streamed_oversized_request.text
    assert client.get("/api/does-not-exist").status_code == 404
    group_a = client.post("/api/groups", json={"name": "Fictional Delivery", "sort_order": 10}).json()
    group_b = client.post("/api/groups", json={"name": "Fictional Governance", "sort_order": 20}).json()
    created = client.post("/api/projects", json={"name": "Fictional Restart Project", "portfolio_group_id": group_a["id"]}).json()
    updated = client.patch(f"/api/projects/{created['id']}", json={
        "status": "Yellow", "priority": "High", "next_action": "Confirm fictional restart",
        "next_action_due": "2026-08-20",
    }).json()
    assert updated["status"] == "Yellow" and updated["priority"] == "High"
    assert Path(updated["folder_path"]).is_dir()
    assert group_b["id"] != group_a["id"]
    with TestClient(
        create_app(settings), base_url="http://127.0.0.1:8765",
        headers={"X-Requested-With": "CHIO-Portfolio-Assistant"},
    ) as restarted:
        persisted = restarted.get(f"/api/projects/{created['id']}").json()
        assert persisted["portfolio_group_id"] == group_a["id"]
        assert persisted["next_action"] == "Confirm fictional restart"


def test_direct_eml_and_msg_preservation_dedup_and_actions(client: TestClient, project, tmp_path: Path):
    first = upload(client, project["id"], "fictional-atlas.eml", make_eml())
    assert first.status_code == 202
    source = first.json()["source"]
    assert process(client, source["id"])["processed"] == 1
    detail = client.get(f"/api/projects/{project['id']}").json()
    assert detail["current_summary"]
    assert len(detail["updates"]) == 1
    assert detail["updates"][0]["citations"]
    assert detail["updates"][0]["citations"][0]["original_filename"] == "fictional-atlas.eml"
    assert detail["updates"][0]["citations"][0]["locator"].startswith("Email body lines")
    assert detail["summary_citations"]
    assert len(detail["action_items"]) == 1
    assert any(item["parent_source_id"] == source["id"] for item in detail["sources"])
    with client.app.state.db.connect() as connection:
        assert connection.execute("SELECT count(*) FROM sources WHERE parent_source_id = ?", (source["id"],)).fetchone()[0] == 1
    duplicate = upload(client, project["id"], "fictional-atlas.eml", make_eml()).json()
    assert duplicate["duplicate"] is True
    process(client, duplicate["source"]["id"])
    with client.app.state.db.connect() as connection:
        assert connection.execute("SELECT count(*) FROM project_updates WHERE project_id = ? AND update_type='knowledge'", (project["id"],)).fetchone()[0] == 1
    msg = upload(client, project["id"], "fictional-readiness.msg", make_msg(tmp_path)).json()["source"]
    assert process(client, msg["id"])["processed"] == 1
    msg_record = next(item for item in client.get(f"/api/projects/{project['id']}").json()["sources"] if item["id"] == msg["id"])
    assert msg_record["processing_state"] == "complete"
    assert client.get(f"/api/sources/{msg['id']}/original").headers["content-disposition"].startswith("attachment")
    assert client.get(f"/api/sources/{msg['id']}/original").headers["content-type"].startswith("application/octet-stream")


def test_capture_removes_renamed_file_when_database_insert_fails(service, project, monkeypatch):
    source_root = Path(project["folder_path"]) / "sources"
    before = {path for path in source_root.rglob("*") if path.is_file()}

    @contextmanager
    def failed_transaction():
        raise sqlite3.OperationalError("fictional database lock timeout")
        yield

    monkeypatch.setattr(service.db, "transaction", failed_transaction)
    with pytest.raises(sqlite3.OperationalError, match="fictional database lock timeout"):
        service.capture_source(
            io.BytesIO(b"Fictional evidence must not become an orphan."),
            "orphan-check.txt", project_id=project["id"],
        )
    after = {path for path in source_root.rglob("*") if path.is_file()}
    assert after == before


def test_transcript_docx_pdf_and_unsupported(client: TestClient, project, tmp_path: Path):
    missing = upload(client, project["id"], "weekly-sync.vtt", b"WEBVTT\n\n00:00.000 --> 00:02.000\nFictional Atlas decision")
    assert missing.status_code == 422
    transcript = upload(
        client, project["id"], "weekly-sync.vtt",
        b"WEBVTT\n\n00:00.000 --> 00:02.000\nFictional Atlas decision approved.",
        meeting_name="Fictional Weekly Sync", meeting_date="2026-08-11", is_transcript=True,
    ).json()["source"]
    assert process(client, transcript["id"])["processed"] == 1
    docx = upload(client, project["id"], "fictional-plan.docx", make_docx(tmp_path)).json()["source"]
    pdf = upload(client, project["id"], "fictional-evidence.pdf", make_pdf(tmp_path, True)).json()["source"]
    assert process(client, docx["id"])["processed"] == 1
    assert process(client, pdf["id"])["processed"] == 1
    scanned = upload(client, project["id"], "fictional-scan.pdf", make_pdf(tmp_path, False)).json()["source"]
    assert process(client, scanned["id"])["unsupported"] == 1
    with client.app.state.db.connect() as connection:
        before_retry = connection.execute(
            "SELECT count(*) FROM source_chunks WHERE source_id = ?", (scanned["id"],)
        ).fetchone()[0]
    assert process(client, scanned["id"])["unsupported"] == 1
    with client.app.state.db.connect() as connection:
        after_retry = connection.execute(
            "SELECT count(*) FROM source_chunks WHERE source_id = ?", (scanned["id"],)
        ).fetchone()[0]
    assert after_retry == before_retry
    unknown = upload(client, project["id"], "fictional-image.png", b"not-an-image").json()["source"]
    assert process(client, unknown["id"])["unsupported"] == 1
    detail = client.get(f"/api/projects/{project['id']}").json()
    transcript_update = next(item for item in detail["updates"] if item["source_id"] == transcript["id"])
    chunk = client.get(f"/api/chunks/{transcript_update['citations'][0]['chunk_id']}").json()
    assert chunk["meeting_name"] == "Fictional Weekly Sync"


def test_ai_outage_and_exact_once_recovery(client: TestClient, project, service):
    service.llm.available = False
    pending = upload(client, project["id"], "outage-note.txt", b"Fictional outage evidence should wait for AI.").json()["source"]
    result = process(client, pending["id"])
    assert result["pending_ai"] == 1
    detail = client.get(f"/api/projects/{project['id']}").json()
    assert next(item for item in detail["sources"] if item["id"] == pending["id"])["processing_state"] == "pending_ai"
    assert not detail["updates"]
    service.llm.available = True
    good = upload(client, project["id"], "independent-note.txt", b"A separate fictional source processes independently.").json()["source"]
    assert process(client, good["id"])["processed"] == 1
    assert client.post("/api/sources/retry-pending").json()["processed"] == 1
    terminal = process(client, pending["id"])
    assert terminal["processed"] == 0 and terminal["already_complete"] == 1
    with client.app.state.db.connect() as connection:
        assert connection.execute("SELECT count(*) FROM project_updates WHERE source_id = ?", (pending["id"],)).fetchone()[0] == 1
        assert connection.execute("SELECT retry_count FROM sources WHERE id = ?", (pending["id"],)).fetchone()[0] == 0


def test_background_worker_claims_once_and_preserves_metadata(settings):
    live_settings = replace(settings, app=replace(settings.app, testing=False))
    with TestClient(
        create_app(live_settings), base_url="http://127.0.0.1:8765",
        headers={"X-Requested-With": "CHIO-Portfolio-Assistant"},
    ) as live_client:
        project = live_client.post("/api/projects", json={"name": "Fictional Worker Project"}).json()
        source = upload(
            live_client, project["id"], "worker-note.txt",
            b"The fictional worker processes this source exactly once.",
        ).json()["source"]
        deadline = time.monotonic() + 5
        state = "captured"
        while time.monotonic() < deadline:
            detail = live_client.get(f"/api/projects/{project['id']}").json()
            state = next(item["processing_state"] for item in detail["sources"] if item["id"] == source["id"])
            if state == "complete":
                break
            time.sleep(0.05)
        assert state == "complete"
        assert live_client.post(f"/api/sources/{source['id']}/retry").json()["processed"] == 0
        with live_client.app.state.db.connect() as connection:
            assert connection.execute(
                "SELECT count(*) FROM project_updates WHERE source_id = ?", (source["id"],)
            ).fetchone()[0] == 1
            metadata = json.loads(connection.execute(
                "SELECT metadata_json FROM sources WHERE id = ?", (source["id"],)
            ).fetchone()[0])
        assert metadata["_extraction_complete"] is True
        assert metadata["evidence_dropped_chunks"] == 0


def test_snow_csv_xlsx_incremental_stale_and_malformed(client: TestClient, service):
    old = "2026-08-10 09:00:00 - Avery Chen (Work note)\nInitial fictional design approved."
    first = client.post("/api/import/servicenow", files={"file": ("snow.csv", snow_csv(old), "text/csv")}).json()
    assert first["tickets_read"] == 1 and first["new_comments_applied"] == 1
    project_id = first["affected_projects"][0]
    client.patch(f"/api/projects/{project_id}", json={"status": "Red", "priority": "Critical"})
    extended = "2026-08-12 09:00:00 - Casey Morgan (Work note)\nSecond new fictional entry.\n2026-08-11 09:00:00 - Avery Chen (Comment)\nFirst new fictional entry.\n" + old
    second = client.post("/api/import/servicenow", files={"file": ("snow.csv", snow_csv(extended, updated="2026-08-12 12:00:00", state="Closed", priority="1"), "text/csv")}).json()
    assert second["new_comments_applied"] == 2
    repeated = client.post("/api/import/servicenow", files={"file": ("snow.csv", snow_csv(extended, updated="2026-08-12 12:00:00", state="Closed", priority="1"), "text/csv")}).json()
    assert repeated["new_comments_applied"] == 0
    stale = client.post("/api/import/servicenow", files={"file": ("snow.csv", snow_csv(old, updated="2026-08-09 12:00:00"), "text/csv")}).json()
    assert stale["new_comments_applied"] == 0
    detail = client.get(f"/api/projects/{project_id}").json()
    assert detail["status"] == "Red" and detail["priority"] == "Critical"
    assert detail["snow_state"] == "Closed" and detail["snow_priority"] == "1"
    with client.app.state.db.connect() as connection:
        comments = connection.execute(
            "SELECT text FROM source_chunks WHERE project_id = ? AND entry_hash IS NOT NULL ORDER BY comment_at", (project_id,)
        ).fetchall()
        assert [row["text"] for row in comments] == ["Initial fictional design approved.", "First new fictional entry.", "Second new fictional entry."]
    xlsx_comment = "2026-08-12 10:30:00 - Jordan Lee (Work note)\nFictional XLSX entry."
    xlsx = client.post("/api/import/servicenow", files={"file": ("snow.xlsx", snow_xlsx(xlsx_comment), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")}).json()
    assert xlsx["new_comments_applied"] == 1
    malformed = client.post("/api/import/servicenow", files={"file": ("bad.csv", snow_csv("Freeform text without a deterministic header", updated="2026-08-13 12:00:00"), "text/csv")}).json()
    assert malformed["review_or_error_count"] == 1
    ragged_stream = io.StringIO()
    ragged_writer = csv.writer(ragged_stream)
    ragged_writer.writerow([
        "Number", "Short description", "Assignment group", "Updated",
        "Comments and Work notes", "State", "Priority",
    ])
    ragged_writer.writerow([
        "REQ0099003", "Fictional Ragged Export", "Fictional Clinical Apps",
        "2026-08-15 12:00:00", old, "In Progress", "2", "unexpected-extra-value",
    ])
    ragged_writer.writerow([
        "REQ0099004", "Fictional Short Export", "Fictional Clinical Apps",
        "2026-08-15 13:00:00", old, "In Progress",
    ])
    ragged = client.post(
        "/api/import/servicenow",
        files={"file": ("ragged.csv", ragged_stream.getvalue().encode(), "text/csv")},
    ).json()
    assert ragged["tickets_read"] == 2 and ragged["review_or_error_count"] == 2
    ragged_ids = set(ragged["review_item_ids"])
    ragged_reviews = [
        item for item in client.get("/api/reviews?status=open").json() if item["id"] in ragged_ids
    ]
    assert len(ragged_reviews) == 2
    assert all("different number of values than the header" in item["reason"] for item in ragged_reviews)
    service.llm.available = False
    outage_comments = "2026-08-14 09:00:00 - Casey Morgan (Work note)\nPending fictional AI entry.\n" + extended
    outage = client.post("/api/import/servicenow", files={"file": ("outage.csv", snow_csv(outage_comments, updated="2026-08-14 12:00:00"), "text/csv")}).json()
    assert outage["pending_ai"] == 1
    with client.app.state.db.connect() as connection:
        before = connection.execute("SELECT snow_comments_cell_hash FROM projects WHERE id = ?", (project_id,)).fetchone()[0]
    service.llm.available = True
    recovered = client.post(
        "/api/import/servicenow",
        files={"file": ("outage-retry.csv", snow_csv(outage_comments, updated="2026-08-14 12:00:00"), "text/csv")},
    ).json()
    assert recovered["new_comments_applied"] == 1 and recovered["tickets_unchanged"] == 0
    service.llm.available = False
    bulk_comments = "2026-08-15 09:00:00 - Avery Chen (Work note)\nBulk retry fictional entry.\n" + outage_comments
    bulk_pending = client.post(
        "/api/import/servicenow",
        files={"file": ("bulk-pending.csv", snow_csv(bulk_comments, updated="2026-08-15 12:00:00"), "text/csv")},
    ).json()
    assert bulk_pending["pending_ai"] == 1
    service.llm.available = True
    assert client.post("/api/sources/retry-pending").json()["processed"] == 1
    with client.app.state.db.connect() as connection:
        after = connection.execute("SELECT snow_comments_cell_hash FROM projects WHERE id = ?", (project_id,)).fetchone()[0]
        assert after != before
        assert connection.execute("SELECT count(*) FROM source_chunks WHERE project_id = ? AND text = 'Pending fictional AI entry.'", (project_id,)).fetchone()[0] == 1
        assert connection.execute("SELECT count(*) FROM source_chunks WHERE project_id = ? AND text = 'Bulk retry fictional entry.'", (project_id,)).fetchone()[0] == 1


def test_project_scoped_chat_and_citation_validation(client: TestClient):
    first = client.post("/api/projects", json={"name": "Fictional North Project"}).json()
    second = client.post("/api/projects", json={"name": "Fictional South Project"}).json()
    north = upload(client, first["id"], "north.txt", b"The fictional launch color is blue and the date is August 20.").json()["source"]
    south = upload(client, second["id"], "south.txt", b"The fictional launch color is red and the date is September 3.").json()["source"]
    process(client, north["id"]); process(client, south["id"])
    answer = client.post(f"/api/projects/{first['id']}/chat", json={"question": "What is the launch color and date?"}).json()
    assert "blue" in answer["answer"] and "red" not in answer["answer"]
    assert answer["claims"][0]["citations"]
    citation = answer["claims"][0]["citations"][0]
    excerpt = client.get(f"/api/chunks/{citation['chunk_id']}").json()
    assert excerpt["project_id"] == first["id"]
    missing = client.post(f"/api/projects/{first['id']}/chat", json={"question": "zzzz-no-such-fact"}).json()
    assert missing["answer"] == "" and missing["claims"] == [] and missing["uncertainty"]


def test_multi_project_review_rule_and_routed_retrieval(client: TestClient):
    atlas = client.post("/api/projects", json={"name": "Fictional Atlas"}).json()
    beacon = client.post("/api/projects", json={"name": "Fictional Beacon"}).json()
    material = b"Fictional Atlas approved its release.\nFictional Beacon needs a new owner."
    captured = client.post(
        "/api/intake/multi-project",
        files={"file": ("cross-workstream.txt", material, "text/plain")},
        data={"meeting_name": "Fictional Portfolio Sync", "meeting_date": "2026-08-12", "is_transcript": "true"},
    ).json()["source"]
    result = process(client, captured["id"])
    assert result["needs_review"] == 1
    assert client.get(f"/api/projects/{atlas['id']}").json()["updates"] == []
    reviews = client.get("/api/reviews?status=open").json()
    review = next(item for item in reviews if item["source_id"] == captured["id"])
    resolved = client.post(f"/api/reviews/{review['id']}/resolve", json={
        "action": "apply", "target_project_id": beacon["id"],
        "rule": review["evidence"][0]["suggested_rule"],
    })
    assert resolved.status_code == 200, resolved.text
    rules = client.get("/api/routing-rules").json()
    assert rules[0]["target_project_id"] == beacon["id"]
    assert rules[0]["created_from_review_id"] == review["id"]
    beacon_detail = client.get(f"/api/projects/{beacon['id']}").json()
    assert len(beacon_detail["updates"]) == 1
    answer = client.post(f"/api/projects/{beacon['id']}/chat", json={"question": "What was approved?"}).json()
    assert answer["claims"]
    second = client.post(
        "/api/intake/multi-project",
        files={"file": ("cross-workstream-followup.txt", b"Follow-up evidence for the recurring fictional workstream.", "text/plain")},
        data={"meeting_name": "Fictional Portfolio Sync", "meeting_date": "2026-08-13", "is_transcript": "true"},
    ).json()["source"]
    process(client, second["id"])
    second_reviews = [item for item in client.get("/api/reviews?status=open").json() if item["source_id"] == second["id"]]
    assert second_reviews and "Matched local routing rule" in second_reviews[0]["reason"]


def test_action_progress_close_protection_and_daily_window(client: TestClient, project, service):
    source = upload(client, project["id"], "action.txt", b"ACTION: Publish fictional guide | OWNER: Me | DUE: 2026-08-20").json()["source"]
    process(client, source["id"])
    action = client.get(f"/api/projects/{project['id']}").json()["action_items"][0]
    invalid_create = client.post(f"/api/projects/{project['id']}/actions", json={
        "description": "Invalid date", "assignee_type": "me", "assignee_value": "Me",
        "due_date": "not-a-date", "state": "open",
    })
    assert invalid_create.status_code == 422
    invalid_update = client.patch(
        f"/api/projects/{project['id']}/actions/{action['id']}", json={"due_date": "2026-99-99"},
    )
    assert invalid_update.status_code == 422
    progress = upload(client, project["id"], "progress.txt", f"PROGRESS [{action['id']}]: Draft is ready for review.".encode()).json()["source"]
    process(client, progress["id"])
    updated = client.get(f"/api/projects/{project['id']}").json()["action_items"][0]
    assert updated["progress_text"] == "Draft is ready for review."
    close = upload(client, project["id"], "close.txt", f"COMPLETE [{action['id']}]".encode()).json()["source"]
    process(client, close["id"])
    assert client.get(f"/api/projects/{project['id']}").json()["action_items"][0]["state"] == "open"
    review = next(item for item in client.get("/api/reviews?status=open").json() if item["kind"] == "action_close")
    client.post(f"/api/reviews/{review['id']}/resolve", json={"action": "apply", "action_item_id": action["id"]})
    assert client.get(f"/api/projects/{project['id']}").json()["action_items"][0]["state"] == "complete"
    assert client.get(f"/api/projects/{project['id']}").json()["status"] == "Green"
    local_now = datetime.now().astimezone()
    start_local = datetime.combine(local_now.date() - timedelta(days=1), datetime.min.time(), tzinfo=local_now.tzinfo)
    inside = (start_local + timedelta(hours=12)).astimezone(timezone.utc).isoformat()
    outside = (start_local - timedelta(seconds=1)).astimezone(timezone.utc).isoformat()
    with client.app.state.db.transaction() as connection:
        inside_id = connection.execute(
            "INSERT INTO project_updates(project_id, source_id, update_type, text, citations_json, created_at) VALUES (?, NULL, 'manual_field', 'Inside prior-day window.', '[]', ?)",
            (project["id"], inside),
        ).lastrowid
        outside_id = connection.execute(
            "INSERT INTO project_updates(project_id, source_id, update_type, text, citations_json, created_at) VALUES (?, NULL, 'manual_field', 'Outside prior-day window.', '[]', ?)",
            (project["id"], outside),
        ).lastrowid
        connection.execute(
            "UPDATE action_items SET updated_at = ? WHERE id = ?", (inside, action["id"])
        )
    first = service.run_daily(date.today())
    original_daily = service.llm.daily
    service.llm.daily = lambda evidence, counts: (_ for _ in ()).throw(
        AssertionError("idempotent daily run called the LLM twice")
    )
    second = service.run_daily(date.today())
    service.llm.daily = original_daily
    assert first["id"] == second["id"]
    cited = second["project_changes"]["cited_update_ids"]
    assert inside_id in cited
    assert outside_id not in cited
    assert second["counts"]["action_changes"] >= 1
    assert any(item["action_item_id"] == action["id"] for item in second["project_changes"]["action_changes"])


def test_incomplete_action_can_be_completed_in_review(client: TestClient, project, service):
    def incomplete_action(summary, evidence):
        citation = [{"source_id": evidence[0]["source_id"], "chunk_id": evidence[0]["chunk_id"]}]
        return {
            "updated_summary": "A fictional action was proposed.",
            "updates": [{"text": "A fictional action was proposed.", "citations": citation}],
            "action_item_operations": [{
                "operation": "create", "description": "Publish fictional guide",
                "assignee_type": "person", "assignee_value": "", "due_date": None,
                "citations": citation, "confidence": 0.5,
            }],
            "project_field_recommendations": [], "needs_review": False, "review_reason": None,
        }

    service.llm.knowledge_update = incomplete_action
    source = upload(client, project["id"], "incomplete-action.txt", b"Publish the fictional guide.").json()["source"]
    assert process(client, source["id"])["processed"] == 1
    review = next(item for item in client.get("/api/reviews?status=open").json() if item["kind"] == "inferred_action")
    resolved = client.post(f"/api/reviews/{review['id']}/resolve", json={
        "action": "apply", "description": "Publish fictional guide", "assignee_type": "person",
        "assignee_value": "Jordan Lee", "due_date": "2026-08-25",
    })
    assert resolved.status_code == 200, resolved.text
    action = client.get(f"/api/projects/{project['id']}").json()["action_items"][0]
    assert action["assignee_value"] == "Jordan Lee" and action["created_by"] == "user"
    assert action["citations"][0]["original_filename"] == "incomplete-action.txt"


def test_direct_project_intake_stops_on_other_project(client: TestClient):
    alpha = client.post("/api/projects", json={"name": "Fictional Alpha Workstream"}).json()
    beta = client.post("/api/projects", json={"name": "Fictional Beta Workstream"}).json()
    source = upload(
        client, alpha["id"], "cross-project-note.txt",
        b"Fictional Beta Workstream approved a separate milestone.",
    ).json()["source"]
    assert process(client, source["id"])["needs_review"] == 1
    assert client.get(f"/api/projects/{alpha['id']}").json()["updates"] == []
    review = next(item for item in client.get("/api/reviews?status=open").json() if item["source_id"] == source["id"])
    assert review["kind"] == "cross_project_evidence"
    assert review["options"] == [{"project_id": beta["id"], "label": "Fictional Beta Workstream"}]
    before_retry = len([
        item for item in client.get("/api/reviews?status=open").json() if item["source_id"] == source["id"]
    ])
    assert process(client, source["id"])["processed"] == 0
    after_retry = len([
        item for item in client.get("/api/reviews?status=open").json() if item["source_id"] == source["id"]
    ])
    assert after_retry == before_retry
    applied = client.post(f"/api/reviews/{review['id']}/resolve", json={
        "action": "apply", "target_project_id": beta["id"],
    })
    assert applied.status_code == 200, applied.text
    beta_detail = client.get(f"/api/projects/{beta['id']}").json()
    assert beta_detail["updates"][0]["citations"][0]["original_filename"] == "cross-project-note.txt"
    alpha_detail = client.get(f"/api/projects/{alpha['id']}").json()
    assert next(item for item in alpha_detail["sources"] if item["id"] == source["id"])["processing_state"] == "complete"
    assert client.get("/api/routing-rules").json() == []


def test_security_size_traversal_and_uncited_rejection(client: TestClient, project, service):
    too_large = upload(client, project["id"], "large.txt", b"x" * (2 * 1024 * 1024 + 1))
    assert too_large.status_code == 422
    assert not any(path.name.endswith("large.txt") for path in Path(project["folder_path"]).rglob("*"))
    captured = upload(client, project["id"], "..\\..\\unsafe.txt", b"Safe fictional content").json()["source"]
    assert ".." not in captured["original_filename"] and Path(captured["original_path"]).is_relative_to(Path(project["folder_path"]))
    original = service.llm.knowledge_update
    service.llm.knowledge_update = lambda summary, evidence: {
        "updated_summary": "Unsupported claim", "updates": [{"text": "Unsupported", "citations": [{"source_id": 999, "chunk_id": 999}]}],
        "action_item_operations": [], "project_field_recommendations": [], "needs_review": False,
    }
    result = process(client, captured["id"])
    assert result["needs_review"] == 1
    assert not client.get(f"/api/projects/{project['id']}").json()["updates"]
    service.llm.knowledge_update = original


def test_evidence_budget_is_the_citation_allowlist(client: TestClient, project, service):
    service.settings = replace(service.settings, llm=replace(service.settings.llm, max_evidence_chars=3000))

    def cite_unsupplied_chunk(summary, evidence):
        assert len(evidence) == 1
        with service.db.connect() as connection:
            dropped = connection.execute(
                "SELECT id, source_id FROM source_chunks WHERE project_id = ? ORDER BY sequence DESC LIMIT 1",
                (project["id"],),
            ).fetchone()
        return {
            "updated_summary": "A fabricated tail claim.",
            "updates": [{"text": "A fabricated tail claim.", "citations": [{"source_id": dropped["source_id"], "chunk_id": dropped["id"]}]}],
            "action_item_operations": [], "project_field_recommendations": [],
            "needs_review": False, "review_reason": None,
        }

    service.llm.knowledge_update = cite_unsupplied_chunk
    material = ("a" * 3000 + "\nThis second chunk was outside the model evidence budget.").encode()
    source = upload(client, project["id"], "bounded.txt", material).json()["source"]
    assert process(client, source["id"])["needs_review"] == 1
    detail = client.get(f"/api/projects/{project['id']}").json()
    assert detail["current_summary"] == ""
    assert next(item for item in detail["sources"] if item["id"] == source["id"])["metadata"]["evidence_dropped_chunks"] == 1


def test_evidence_budget_skips_oversized_middle_chunk(service):
    service.settings = replace(service.settings, llm=replace(service.settings.llm, max_evidence_chars=10))
    evidence = [
        {"source_id": 1, "chunk_id": 1, "text": "123456"},
        {"source_id": 1, "chunk_id": 2, "text": "too-large-for-remaining"},
        {"source_id": 1, "chunk_id": 3, "text": "abc"},
    ]
    bounded, dropped = service._bounded_evidence(evidence)
    assert [item["chunk_id"] for item in bounded] == [1, 3]
    assert dropped == 1


def test_uncertainty_cannot_cover_an_uncited_factual_answer(client: TestClient, project, service):
    source = upload(client, project["id"], "grounding.txt", b"The fictional launch date is August 20.").json()["source"]
    process(client, source["id"])
    service.llm.chat = lambda question, summary, evidence: {
        "answer": "The launch date moved to September 3.",
        "claims": [],
        "uncertainty": "Some details may be incomplete.",
    }
    with pytest.raises(LlmContractError, match="must contain claim citations"):
        service.ask_project(project["id"], "When is launch?")


def test_literal_wildcard_search_and_attachment_restart_recovery(client: TestClient, project, service):
    source = upload(client, project["id"], "parent.txt", b"Ordinary fictional evidence without wildcard characters.").json()["source"]
    process(client, source["id"])
    assert service.db.search_chunks(project["id"], "%") == []

    attachment_path = Path(project["folder_path"]) / "recovered-attachment.txt"
    attachment_path.write_text("Recovered fictional attachment evidence.", encoding="utf-8")
    digest = hashlib.sha256(attachment_path.read_bytes()).hexdigest()
    with service.db.transaction() as connection:
        child_id = connection.execute(
            """INSERT INTO sources(project_id, parent_source_id, source_type, native_id, sha256,
               original_filename, original_path, processing_state, created_at)
               VALUES (?, ?, 'txt', ?, ?, 'recovered-attachment.txt', ?, 'processing', ?)""",
            (project["id"], source["id"], f"attachment:{source['id']}:99", digest, str(attachment_path), datetime.now(timezone.utc).isoformat()),
        ).lastrowid
    assert service.recover_interrupted() >= 1
    with service.db.connect() as connection:
        recovered = connection.execute("SELECT processing_state FROM sources WHERE id = ?", (child_id,)).fetchone()
        chunks = connection.execute("SELECT count(*) FROM source_chunks WHERE source_id = ?", (child_id,)).fetchone()[0]
    assert recovered["processing_state"] == "complete"
    assert chunks == 1
